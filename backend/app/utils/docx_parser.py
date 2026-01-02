"""
Word文档解析模块
"""
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


class DocxParser:
    """Word文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化解析器
        
        Args:
            file_path: Word文档路径
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        self.file_path = file_path
        self.document = Document(file_path)
    
    def get_title(self) -> str:
        """
        获取文档标题（通常是第一个段落或标题样式）
        
        Returns:
            文档标题
        """
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                return paragraph.text.strip()
        return ""
    
    def get_all_text(self) -> str:
        """
        获取文档所有文本
        
        Returns:
            全文文本
        """
        return "\n".join([para.text for para in self.document.paragraphs])
    
    def get_paragraphs(self) -> List[Dict[str, Any]]:
        """
        获取所有段落及其属性
        
        Returns:
            段落列表，每个段落包含文本和格式信息
        """
        paragraphs = []
        
        for idx, para in enumerate(self.document.paragraphs):
            para_info = {
                "index": idx,
                "text": para.text,
                "style": para.style.name if para.style else None,
                "alignment": para.alignment,
                "is_empty": len(para.text.strip()) == 0
            }
            
            # 获取段落格式
            if para.paragraph_format:
                para_info.update({
                    "line_spacing": para.paragraph_format.line_spacing,
                    "space_before": para.paragraph_format.space_before,
                    "space_after": para.paragraph_format.space_after,
                    "left_indent": para.paragraph_format.left_indent,
                    "first_line_indent": para.paragraph_format.first_line_indent
                })
            
            paragraphs.append(para_info)
        
        return paragraphs
    
    def get_headings(self) -> List[Dict[str, Any]]:
        """
        获取所有标题
        
        Returns:
            标题列表
        """
        headings = []
        
        for idx, para in enumerate(self.document.paragraphs):
            if para.style and 'Heading' in para.style.name:
                level = 0
                if 'Heading 1' in para.style.name:
                    level = 1
                elif 'Heading 2' in para.style.name:
                    level = 2
                elif 'Heading 3' in para.style.name:
                    level = 3
                
                headings.append({
                    "index": idx,
                    "level": level,
                    "text": para.text,
                    "style": para.style.name
                })
        
        return headings
    
    def get_tables(self) -> List[Dict[str, Any]]:
        """
        获取所有表格
        
        Returns:
            表格列表
        """
        tables = []
        
        for idx, table in enumerate(self.document.tables):
            table_info = {
                "index": idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": []
            }
            
            # 提取表格数据
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_info["data"].append(row_data)
            
            tables.append(table_info)
        
        return tables
    
    def get_images(self) -> List[Dict[str, Any]]:
        """
        获取所有图片信息
        
        Returns:
            图片列表
        """
        images = []
        
        for idx, rel in enumerate(self.document.part.rels.values()):
            if "image" in rel.target_ref:
                images.append({
                    "index": idx,
                    "filename": os.path.basename(rel.target_ref),
                    "type": rel.target_ref.split('.')[-1]
                })
        
        return images
    
    def get_page_setup(self) -> Dict[str, Any]:
        """
        获取页面设置信息
        
        Returns:
            页面设置字典
        """
        section = self.document.sections[0]
        
        return {
            "page_width": section.page_width,
            "page_height": section.page_height,
            "left_margin": section.left_margin,
            "right_margin": section.right_margin,
            "top_margin": section.top_margin,
            "bottom_margin": section.bottom_margin,
            "orientation": "Portrait" if section.page_width < section.page_height else "Landscape"
        }
    
    def get_fonts(self) -> List[Dict[str, Any]]:
        """
        获取文档中使用的字体信息
        
        Returns:
            字体列表
        """
        fonts = []
        seen_fonts = set()
        
        for para in self.document.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name not in seen_fonts:
                    fonts.append({
                        "name": run.font.name,
                        "size": run.font.size.pt if run.font.size else None,
                        "bold": run.font.bold,
                        "italic": run.font.italic
                    })
                    seen_fonts.add(run.font.name)
        
        return fonts
    
    def extract_keywords(self) -> List[str]:
        """
        提取关键词（简单实现：提取前几段的重要词汇）
        
        Returns:
            关键词列表
        """
        # 简单实现：提取前3段的非空词汇
        keywords = []
        count = 0
        
        for para in self.document.paragraphs:
            if para.text.strip() and count < 3:
                words = para.text.split()
                keywords.extend([w for w in words if len(w) > 2])
                count += 1
        
        # 去重并返回前10个
        return list(set(keywords))[:10]
    
    def get_summary(self) -> Dict[str, Any]:
        """
        获取文档摘要信息
        
        Returns:
            摘要信息字典
        """
        return {
            "file_path": self.file_path,
            "title": self.get_title(),
            "paragraph_count": len(self.document.paragraphs),
            "table_count": len(self.document.tables),
            "image_count": len(self.get_images()),
            "heading_count": len(self.get_headings()),
            "total_words": len(self.get_all_text().split()),
            "page_setup": self.get_page_setup()
        }
